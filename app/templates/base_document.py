"""
Base Document
-------------

Lớp cơ sở cho toàn bộ Template Engine.

Chịu trách nhiệm:

- Khởi tạo Document.
- Thiết lập Page.
- Thiết lập Font.
- Điều phối các Component.
"""

from docx import Document


from app.templates.components.header import HeaderComponent
from app.templates.components.recipient import RecipientComponent
from app.templates.components.signature import SignatureComponent

from app.templates.styles import (
    ALIGN_CENTER,
    ALIGN_JUSTIFY,
    ALIGN_LEFT,
    ALIGN_RIGHT,
    apply_font,
    apply_normal_style,
    apply_page_style,
    apply_paragraph_style,
)


class BaseDocument:
    """
    Base class của toàn bộ Template.
    """

    def __init__(self):

        self.doc = Document()

        self._initialize()

    # =====================================================
    # INITIALIZE
    # =====================================================

    def _initialize(self) -> None:
        """
        Khởi tạo Document.
        """

        self.setup_page()

        self.setup_font()

    # =====================================================
    # PAGE
    # =====================================================

    def setup_page(self) -> None:
        """
        Thiết lập khổ giấy.
        """

        apply_page_style(self.doc)

    # =====================================================
    # FONT
    # =====================================================

    def setup_font(self) -> None:
        """
        Thiết lập font mặc định.
        """

        apply_normal_style(self.doc)

    # =====================================================
    # PARAGRAPH
    # =====================================================

    def paragraph(
        self,
        text: str = "",
        *,
        bold: bool = False,
        italic: bool = False,
        align=ALIGN_JUSTIFY,
        size: int = 13,
    ):
        """
        Thêm một Paragraph chuẩn.
        """

        paragraph = self.doc.add_paragraph()

        apply_paragraph_style(
            paragraph,
            align=align,
        )

        run = paragraph.add_run(text)

        apply_font(
            run,
            size=size,
            bold=bold,
            italic=italic,
        )

        return paragraph

    # =====================================================
    # TEXT
    # =====================================================

    def add_text(
        self,
        text: str,
        *,
        align=ALIGN_JUSTIFY,
    ) -> None:
        """
        Thêm nhiều đoạn văn.
        """

        for line in text.splitlines():

            line = line.strip()

            if not line:
                continue

            self.paragraph(
                text=line,
                align=align,
            )

    # =====================================================
    # HEADING
    # =====================================================

    def add_heading(
        self,
        text: str,
        level: int = 1,
        align=ALIGN_CENTER,
    ):
        """
        Thêm Heading.
        """

        heading_style = {
            1: (15, True),
            2: (14, True),
            3: (13, True),
        }

        size, bold = heading_style.get(
            level,
            (13, False),
        )

        paragraph = self.doc.add_paragraph()

        apply_paragraph_style(
            paragraph,
            align=align,
            indent=False,
        )

        run = paragraph.add_run(text)

        apply_font(
            run,
            size=size,
            bold=bold,
        )

        return paragraph

    # =====================================================
    # DOCUMENT API
    # =====================================================

    def create_content(
        self,
        text: str,
    ) -> None:
        """
        Sinh nội dung văn bản.
        """

        self.add_text(text)

    def blank(
        self,
        lines: int = 1,
    ) -> None:
        """
        Chèn dòng trắng.
        """

        for _ in range(lines):
            self.doc.add_paragraph()


    # =====================================================
    # HEADER
    # =====================================================

    def create_header(self, agency, unit, location, date_text, number=""):
        HeaderComponent(self.doc).render(
            agency=agency,
            unit=unit,
            location=location,
            date_text=date_text,
            number=number,
        )
        self.blank()

    # =====================================================
    # TITLE
    # =====================================================

    def create_title(self, title, subtitle=""):
        self.add_heading(title.upper(), level=1)
        if subtitle:
            self.add_heading(subtitle, level=3)
        self.blank()

    # =====================================================
    # SIGNATURE
    # =====================================================

    def create_signature(self, position, signer):
        SignatureComponent(self.doc).render(
            position=position,
            signer=signer,
        )

    # =====================================================
    # RECIPIENT
    # =====================================================

    def add_recipient(self, recipients):
        RecipientComponent(self.doc).build(recipients)

    # =====================================================
    # FOOTER
    # =====================================================

    def add_footer(self, text=""):
        if not text:
            return
        self.blank()
        p=self.doc.add_paragraph()
        apply_paragraph_style(p, align=ALIGN_CENTER, indent=False)
        run=p.add_run(text)
        apply_font(run, italic=True, size=11)

    # =====================================================
    # TABLE
    # =====================================================

    def add_table(self, headers, rows):
        table=self.doc.add_table(rows=1, cols=len(headers))
        table.style="Table Grid"
        hdr=table.rows[0].cells
        for i,h in enumerate(headers):
            p=hdr[i].paragraphs[0]
            apply_paragraph_style(p, align=ALIGN_CENTER, indent=False)
            r=p.add_run(str(h))
            apply_font(r, bold=True)
        for row in rows:
            cells=table.add_row().cells
            for i,v in enumerate(row):
                p=cells[i].paragraphs[0]
                apply_paragraph_style(p, align=ALIGN_LEFT, indent=False)
                r=p.add_run(str(v))
                apply_font(r)
        return table

    def save(
        self,
        path: str,
    ) -> None:
        """
        Lưu file Word.
        """

        self.doc.save(path)