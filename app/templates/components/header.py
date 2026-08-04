"""
Header Component
----------------

Sinh Header chuẩn theo Nghị định 30/2020/NĐ-CP.
"""

from docx.document import Document
from docx.shared import Cm

from app.templates.styles import (
    ALIGN_CENTER,
    apply_font,
    apply_paragraph_style,
    HEADER_LEFT_WIDTH,
    HEADER_RIGHT_WIDTH,
)


class HeaderComponent:
    """
    Component tạo Header văn bản.
    """

    def __init__(
        self,
        document: Document,
    ):

        self.document = document

    # =====================================================
    # PUBLIC
    # =====================================================

    def render(
        self,
        *,
        agency: str,
        unit: str,
        location: str,
        date_text: str,
        number: str = "",
    ) -> None:
        """
        Sinh Header văn bản.
        """

        table = self.document.add_table(
            rows=1,
            cols=2,
        )

        table.autofit = False

        table.columns[0].width = HEADER_LEFT_WIDTH

        table.columns[1].width = HEADER_RIGHT_WIDTH

        left = table.cell(0, 0)

        right = table.cell(0, 1)

        self._render_left(
            left,
            agency,
            unit,
            number,
        )

        self._render_right(
            right,
            location,
            date_text,
        )

        self.document.add_paragraph()

    # =====================================================
    # PRIVATE
    # =====================================================

    def _render_left(
        self,
        cell,
        agency: str,
        unit: str,
        number: str,
    ) -> None:

        p = cell.paragraphs[0]

        apply_paragraph_style(
            p,
            align=ALIGN_CENTER,
            indent=False,
        )

        run = p.add_run(
            agency + "\n"
        )

        apply_font(
            run,
            bold=True,
        )

        run = p.add_run(unit)

        apply_font(
            run,
            bold=True,
        )

        p = cell.add_paragraph()

        apply_paragraph_style(
            p,
            align=ALIGN_CENTER,
            indent=False,
        )

        p.add_run("_______________")

        p = cell.add_paragraph()

        apply_paragraph_style(
            p,
            align=ALIGN_CENTER,
            indent=False,
        )

        run = p.add_run(
            f"Số: {number}"
        )

        apply_font(run)

    def _render_right(
        self,
        cell,
        location: str,
        date_text: str,
    ) -> None:

        p = cell.paragraphs[0]

        apply_paragraph_style(
            p,
            align=ALIGN_CENTER,
            indent=False,
        )

        run = p.add_run(
            "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n"
        )

        apply_font(
            run,
            bold=True,
        )

        run = p.add_run(
            "Độc lập - Tự do - Hạnh phúc"
        )

        apply_font(
            run,
            bold=True,
        )

        p = cell.add_paragraph()

        apply_paragraph_style(
            p,
            align=ALIGN_CENTER,
            indent=False,
        )

        p.add_run("_______________")

        p = cell.add_paragraph()

        apply_paragraph_style(
            p,
            align=ALIGN_CENTER,
            indent=False,
        )

        run = p.add_run(
            f"{location}, {date_text}"
        )

        apply_font(
            run,
            italic=True,
        )