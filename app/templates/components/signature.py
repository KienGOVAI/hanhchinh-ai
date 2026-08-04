"""
Signature Component
-------------------

Component tạo phần chữ ký của văn bản.
"""

from docx.document import Document

from app.templates.styles import (
    ALIGN_RIGHT,
    apply_font,
    apply_paragraph_style,
)


class SignatureComponent:
    """
    Component tạo phần chữ ký.
    """

    def __init__(
        self,
        document: Document,
    ):

        self.document = document

    # =====================================================
    # PUBLIC
    # =====================================================

    def render(
        self,
        *,
        position: str,
        signer: str,
    ) -> None:
        """
        Sinh phần chữ ký.
        """

        paragraph = self.document.add_paragraph()

        apply_paragraph_style(
            paragraph,
            align=ALIGN_RIGHT,
            indent=False,
        )

        run = paragraph.add_run(position)

        apply_font(
            run,
            bold=True,
        )

        # Chừa khoảng trống để ký
        for _ in range(3):
            self.document.add_paragraph()

        paragraph = self.document.add_paragraph()

        apply_paragraph_style(
            paragraph,
            align=ALIGN_RIGHT,
            indent=False,
        )

        run = paragraph.add_run(signer)

        apply_font(
            run,
            bold=True,
        )