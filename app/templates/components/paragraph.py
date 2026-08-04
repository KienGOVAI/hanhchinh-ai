"""
Paragraph Component
-------------------

Component tạo Paragraph chuẩn cho văn bản hành chính.
"""

from docx.document import Document

from app.templates.styles import (
    ALIGN_CENTER,
    ALIGN_JUSTIFY,
    apply_font,
    apply_paragraph_style,
)


class ParagraphComponent:
    """
    Component xử lý Paragraph.
    """

    def __init__(
        self,
        document: Document,
    ):

        self.document = document

    # =====================================================
    # PUBLIC
    # =====================================================

    def add(
        self,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        align=ALIGN_JUSTIFY,
        size: int = 13,
        indent: bool = True,
    ):
        """
        Thêm một Paragraph.
        """

        paragraph = self.document.add_paragraph()

        apply_paragraph_style(
            paragraph,
            align=align,
            indent=indent,
        )

        run = paragraph.add_run(text)

        apply_font(
            run,
            size=size,
            bold=bold,
            italic=italic,
        )

        return paragraph

    def add_text(
        self,
        text: str,
        *,
        align=ALIGN_JUSTIFY,
    ) -> None:
        """
        Thêm nhiều Paragraph từ chuỗi nhiều dòng.
        """

        for line in text.splitlines():

            line = line.strip()

            if not line:
                continue

            self.add(
                line,
                align=align,
            )

    def add_heading(
        self,
        text: str,
        *,
        level: int = 1,
        align=ALIGN_CENTER,
    ):
        """
        Thêm tiêu đề.
        """

        heading_style = {
            1: (15, True),
            2: (14, True),
            3: (13, True),
        }

        size, bold = heading_style.get(
            level,
            (13, False),
        )

        return self.add(
            text=text,
            bold=bold,
            size=size,
            align=align,
            indent=False,
        )

    def blank(
        self,
        lines: int = 1,
    ) -> None:
        """
        Chèn dòng trắng.
        """

        for _ in range(lines):

            self.document.add_paragraph()