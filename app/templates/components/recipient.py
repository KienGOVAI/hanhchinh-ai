"""
Recipient Component
-------------------

Component tạo phần "Nơi nhận" của văn bản hành chính.
"""

from docx.document import Document

from app.templates.styles import (
    ALIGN_LEFT,
    apply_font,
    apply_paragraph_style,
)


class RecipientComponent:
    """
    Component tạo mục Nơi nhận.
    """

    def __init__(
        self,
        document: Document,
    ):

        self.document = document

    # =====================================================
    # PUBLIC
    # =====================================================

    def render(
        self,
        recipients: list[str],
    ) -> None:
        """
        Sinh mục Nơi nhận.
        """

        if not recipients:
            return

        title = self.document.add_paragraph()

        apply_paragraph_style(
            title,
            align=ALIGN_LEFT,
            indent=False,
        )

        run = title.add_run("Nơi nhận:")

        apply_font(
            run,
            bold=True,
            italic=True,
        )

        for item in recipients:

            paragraph = self.document.add_paragraph()

            apply_paragraph_style(
                paragraph,
                align=ALIGN_LEFT,
                indent=False,
            )

            run = paragraph.add_run(
                f"- {item}"
            )

            apply_font(run)