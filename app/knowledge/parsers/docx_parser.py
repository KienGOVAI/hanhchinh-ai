"""
DOCX Parser
-----------

Parser Microsoft Word DOCX.
"""

from pathlib import Path

from docx import Document

from app.knowledge.models.parsed_document import (
    ParsedDocument,
    ParsedPage,
)

from app.knowledge.parsers.base_parser import (
    BaseParser,
    EmptyDocumentError,
    FileReadError,
    InvalidDocumentError,
)


class DOCXParser(BaseParser):
    """
    Parser DOCX.
    """

    def parse(
        self,
        file_path: str | Path,
    ) -> ParsedDocument:

        path = Path(file_path)

        if not path.exists():
            raise FileReadError(
                f"Không tìm thấy file: {path}"
            )

        try:
            document = Document(path)
        except Exception as exc:
            raise InvalidDocumentError(
                f"Không thể mở DOCX: {path}"
            ) from exc

        blocks: list[str] = []

        # -----------------------------------------
        # PARAGRAPHS
        # -----------------------------------------

        for paragraph in document.paragraphs:

            text = self._normalize(
                paragraph.text
            )

            if text:
                blocks.append(text)

        # -----------------------------------------
        # TABLES
        # -----------------------------------------

        for table in document.tables:

            for row in table.rows:

                cells = [
                    self._normalize(
                        cell.text
                    )
                    for cell in row.cells
                ]

                cells = [
                    cell
                    for cell in cells
                    if cell
                ]

                if cells:
                    blocks.append(
                        " | ".join(cells)
                    )

        full_text = "\n".join(
            blocks
        ).strip()

        if not full_text:
            raise EmptyDocumentError(
                f"DOCX không chứa nội dung: {path}"
            )

        page = ParsedPage(
            page_number=1,
            text=full_text,
        )

        return ParsedDocument(
            text=full_text,
            source=path.name,
            pages=[page],
            metadata={
                "format": "docx",
                "filename": path.name,
            },
        )

    @staticmethod
    def _normalize(text: str) -> str:
        return " ".join(
            text.split()
        ).strip()